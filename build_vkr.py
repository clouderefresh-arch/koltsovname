"""Сборка итогового .docx файла ВКР."""
from docx import Document
from docx.shared import Pt, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from docx_styles import (
    configure_page, configure_base_styles, set_first_page_no_number,
    add_page_number_footer, add_para, add_struct_heading, add_chapter_heading,
    add_para_heading, add_caption, add_source_line, add_table, add_pagebreak,
    set_run_font,
)

from vkr_content import CONTENT


OUTPUT = "/workspace/output/VKR_KupiShuz_revised.docx"


def add_title_page(doc: Document):
    add_para(doc, "АККРЕДИТОВАННОЕ ОБРАЗОВАТЕЛЬНОЕ ЧАСТНОЕ УЧРЕЖДЕНИЕ",
             bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    add_para(doc, "ВЫСШЕГО ОБРАЗОВАНИЯ",
             bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    add_para(doc, "«МОСКОВСКИЙ ФИНАНСОВО-ЮРИДИЧЕСКИЙ УНИВЕРСИТЕТ МФЮА»",
             bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, space_after=12)
    add_para(doc, "Кафедра менеджмента",
             align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, space_after=24)
    add_para(doc, "Специальность 38.02.04 Коммерция (по отраслям)",
             align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, space_after=36)

    add_para(doc, "ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА",
             bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    add_para(doc, "(дипломная работа)",
             align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, space_after=24)
    add_para(doc, "на тему:",
             align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    add_para(doc,
             "«Разработка бизнес-плана по производству товара (на материалах ООО „КУПИШУЗ“)»",
             bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, space_after=48)

    add_para(doc, "Выполнил(а) обучающийся(аяся) ___ курса",
             align=WD_ALIGN_PARAGRAPH.LEFT, indent=False)
    add_para(doc, "_________________________________________ (Ф. И. О.)",
             align=WD_ALIGN_PARAGRAPH.LEFT, indent=False, space_after=6)
    add_para(doc, "Подпись _____________",
             align=WD_ALIGN_PARAGRAPH.LEFT, indent=False, space_after=18)
    add_para(doc, "Научный руководитель",
             align=WD_ALIGN_PARAGRAPH.LEFT, indent=False)
    add_para(doc, "_________________________________________ (Ф. И. О., должность, учёная степень)",
             align=WD_ALIGN_PARAGRAPH.LEFT, indent=False, space_after=6)
    add_para(doc, "Подпись _____________",
             align=WD_ALIGN_PARAGRAPH.LEFT, indent=False, space_after=18)

    add_para(doc, "Дата защиты «___» ________________ 20___ г.",
             align=WD_ALIGN_PARAGRAPH.LEFT, indent=False)
    add_para(doc, "Оценка «_____________»",
             align=WD_ALIGN_PARAGRAPH.LEFT, indent=False, space_after=24)
    add_para(doc, "Москва 2026", bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)


def add_placeholder_page(doc, text):
    add_para(doc, text.upper(), bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, space_before=36)
    add_para(doc, "(лист резервируется для распечатки и подшивки в работу)",
             align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, italic=True)


def add_toc(doc):
    """Поле {TOC} для автоматического оглавления Word."""
    p = doc.add_paragraph()
    run = p.add_run()
    set_run_font(run, 14)
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Оглавление будет обновлено автоматически в Microsoft Word (Ctrl+A → F9)."
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), "Times New Roman")
    rPr.append(rFonts)
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instr)
    run._r.append(fldChar2)
    run._r.append(text)
    run._r.append(fldChar3)


def add_figure_placeholder(doc, caption, desc=None):
    add_para(doc, "[Место для рисунка]", align=WD_ALIGN_PARAGRAPH.CENTER,
             indent=False, italic=True, space_before=6, space_after=6)
    add_para(doc, caption, align=WD_ALIGN_PARAGRAPH.CENTER,
             indent=False, space_after=0)
    if desc:
        add_source_line(doc, desc)


def add_formula(doc, text, number):
    add_para(doc, f"{text}    {number}", align=WD_ALIGN_PARAGRAPH.CENTER,
             indent=False, space_before=6, space_after=6)


def add_where_block(doc, items):
    add_para(doc, "где:", align=WD_ALIGN_PARAGRAPH.LEFT, indent=False,
             space_after=0)
    for it in items:
        add_para(doc, it, align=WD_ALIGN_PARAGRAPH.LEFT, indent=False,
                 first_line_indent_cm=0, space_after=0)


def add_dash_list(doc, items):
    for it in items:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.first_line_indent = Cm(1.25)
        pf.line_spacing = 1.5
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run("– " + it)
        set_run_font(run, 14)


def add_ref_list(doc, items, start_num):
    n = start_num
    for it in items:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.first_line_indent = Cm(1.25)
        pf.line_spacing = 1.5
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(f"{n}. {it}")
        set_run_font(run, 14)
        n += 1
    return n


def render():
    doc = Document()
    configure_base_styles(doc)

    section = doc.sections[0]
    configure_page(section)
    add_page_number_footer(section)
    set_first_page_no_number(section)

    ref_counter = 1

    for item in CONTENT:
        t = item.get("type")
        if t == "title_page":
            add_title_page(doc)
        elif t == "page_break":
            add_pagebreak(doc)
        elif t == "placeholder_page":
            add_placeholder_page(doc, item["text"])
        elif t == "struct_heading":
            add_struct_heading(doc, item["text"])
        elif t == "toc":
            add_toc(doc)
        elif t == "chapter_heading":
            add_chapter_heading(doc, item["text"])
        elif t == "para_heading":
            add_para_heading(doc, item["text"])
        elif t == "subsection_heading":
            add_para(doc, item["text"], bold=True,
                     align=WD_ALIGN_PARAGRAPH.CENTER, indent=False,
                     space_before=12, space_after=6, keep_with_next=True)
        elif t == "appendix_heading":
            add_para(doc, item["text"], bold=True,
                     align=WD_ALIGN_PARAGRAPH.CENTER, indent=False,
                     space_before=12, space_after=12, keep_with_next=True)
        elif t == "p":
            add_para(doc, item["text"])
        elif t == "list":
            add_dash_list(doc, item["items"])
        elif t == "ref_list":
            ref_counter = add_ref_list(doc, item["items"], ref_counter)
        elif t == "table":
            add_table(doc, item["header"], item["rows"],
                      col_widths_cm=item.get("col_widths_cm"))
        elif t == "caption_table":
            add_caption(doc, item["text"], align=WD_ALIGN_PARAGRAPH.LEFT)
        elif t == "source":
            add_source_line(doc, item["text"])
        elif t == "figure":
            add_figure_placeholder(doc, item["caption"], item.get("desc"))
        elif t == "figure_placeholder":
            add_figure_placeholder(doc, item["caption"])
        elif t == "formula":
            add_formula(doc, item["text"], item["number"])
        elif t == "where":
            add_where_block(doc, item["items"])
        else:
            print("unknown item type:", t)

    import os
    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
    doc.save(OUTPUT)
    print("Saved:", OUTPUT)


if __name__ == "__main__":
    render()
