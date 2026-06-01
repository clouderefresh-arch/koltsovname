"""Утилиты для оформления ВКР согласно методическим рекомендациям МФЮА."""
from docx import Document
from docx.shared import Pt, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


FONT_NAME = "Times New Roman"


def set_run_font(run, size_pt: float = 14, bold: bool = False, italic: bool = False):
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), FONT_NAME)


def configure_page(section, top_mm=20, bottom_mm=20, left_mm=30, right_mm=10):
    section.top_margin = Mm(top_mm)
    section.bottom_margin = Mm(bottom_mm)
    section.left_margin = Mm(left_mm)
    section.right_margin = Mm(right_mm)
    section.page_height = Mm(297)
    section.page_width = Mm(210)


def add_page_number_footer(section):
    footer = section.footer
    # Remove any existing default text
    for p in list(footer.paragraphs):
        p._element.getparent().remove(p._element)
    p = footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    run = p.add_run()
    set_run_font(run, 14)
    # Field: PAGE
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE \\* MERGEFORMAT"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)


def set_first_page_no_number(section):
    """Заставляем первый лист (титульный) не печатать номер страницы."""
    sectPr = section._sectPr
    titlePg = sectPr.find(qn("w:titlePg"))
    if titlePg is None:
        titlePg = OxmlElement("w:titlePg")
        sectPr.append(titlePg)
    # Empty first-page footer
    first_footer = section.first_page_footer
    for p in list(first_footer.paragraphs):
        p._element.getparent().remove(p._element)
    p = first_footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def configure_base_styles(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(14)
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Cm(1.25)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # Шрифт для кириллицы
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONT_NAME)


def add_para(doc, text, *, bold=False, italic=False, align=None, indent=True,
             first_line_indent_cm=1.25, space_before=0, space_after=0,
             keep_with_next=False, size=14, line_spacing=1.5):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.first_line_indent = Cm(first_line_indent_cm) if indent else Cm(0)
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.keep_with_next = keep_with_next
    if align is not None:
        p.alignment = align
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run_font(run, size, bold=bold, italic=italic)
    return p


def add_struct_heading(doc, text):
    """Заголовок структурного элемента (ОГЛАВЛЕНИЕ, ВВЕДЕНИЕ ...) – по центру, прописными, жирный."""
    return add_para(doc, text.upper(), bold=True,
                    align=WD_ALIGN_PARAGRAPH.CENTER, indent=False,
                    space_after=12, keep_with_next=True)


def add_chapter_heading(doc, text):
    """Заголовок главы – по центру, прописными, жирный."""
    return add_para(doc, text.upper(), bold=True,
                    align=WD_ALIGN_PARAGRAPH.CENTER, indent=False,
                    space_before=0, space_after=12, keep_with_next=True)


def add_para_heading(doc, text):
    """Заголовок параграфа – с абзацного отступа, жирный."""
    return add_para(doc, text, bold=True,
                    align=WD_ALIGN_PARAGRAPH.LEFT, indent=True,
                    space_before=12, space_after=6, keep_with_next=True)


def add_caption(doc, text, *, align=WD_ALIGN_PARAGRAPH.LEFT, indent=False):
    """Подпись таблицы (слева) или рисунка (по центру)."""
    return add_para(doc, text, align=align, indent=indent,
                    space_before=6, space_after=0, keep_with_next=True)


def add_source_line(doc, text):
    """Строка «Источник: …» под таблицей – без отступа, выравнивание по ширине, 12pt."""
    return add_para(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=False,
                    space_before=0, space_after=6, size=12, line_spacing=1.0)


def add_table(doc, header, rows, *, first_col_width_cm=None, col_widths_cm=None):
    """Создаёт таблицу с шапкой и строками. Шрифт 12pt, межстрочный 1.0, без выделений."""
    n_cols = len(header)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "Table Grid"
    table.autofit = False
    # Width columns
    if col_widths_cm is None:
        total_cm = 16.0
        col_widths_cm = [total_cm / n_cols] * n_cols
    if first_col_width_cm is not None:
        col_widths_cm = list(col_widths_cm)
        col_widths_cm[0] = first_col_width_cm
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Cm(col_widths_cm[idx])

    def fill_cell(cell, text, *, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in list(cell.paragraphs):
            p._element.getparent().remove(p._element)
        p = cell.add_paragraph()
        p.alignment = align
        pf = p.paragraph_format
        pf.first_line_indent = Cm(0)
        pf.line_spacing = 1.0
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        run = p.add_run(str(text))
        set_run_font(run, 12, bold=bold)

    for idx, h in enumerate(header):
        fill_cell(table.rows[0].cells[idx], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            fill_cell(table.rows[r_idx + 1].cells[c_idx], value, align=align)
    return table


def add_pagebreak(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    set_run_font(run, 14)
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)
